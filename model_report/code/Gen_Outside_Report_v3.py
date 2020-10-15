import codecs
import json
import os
import seaborn as sns
from sklearn.model_selection import train_test_split
from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from sklearn import metrics
from matplotlib import pylab
import matplotlib.pyplot as plt
import numpy as np
from al_plot import *
from al_statis import *
from matplotlib.font_manager import FontProperties
from al_data_tool import *
from al_metrics import *
import math
import warnings
import matplotlib
from datetime import datetime

matplotlib.rcParams['font.sans-serif'] = ['SimHei']
matplotlib.rcParams['font.family']='sans-serif'
plt.rcParams['axes.unicode_minus']=False
warnings.filterwarnings('ignore')

def get_data(df):
    target = 'y_label'
    x_train, x_test, y_train, y_test = train_test_split(df, df[target], test_size=0.3, stratify=df[target])
    print(abs(y_train.sum()/len(y_train)-y_test.sum()/len(y_test)))
    x_train['type'] = 'train'
    x_test['type'] = 'test'
    df = pd.concat([x_train, x_test], ignore_index=True)
    df.index = range(len(df))
    return df

def merge(df):
    oot = df.loc[df['triger_date']>'2020-03']
    train = df.loc[df['triger_date']<'2020-03']
    train = get_data(train)
    df = pd.concat([oot, train], ignore_index=True)
    df.index = range(len(df))
    return df

class ML_External_Report(object):
    def __init__(self, df_src,df_src_month,conf_dict):
        self.df_src = df_src
        self.conf_dict = conf_dict
        self.df_src_month =df_src_month
        self.df_src['type'] = self.df_src['type'].map(lambda s: s.lower())
        self.splitted_type_order = ['train', 'test', 'oot']
        if "oot" not in set(df_src['type']):
            self.splitted_type_order = ['train', 'test']
        self.method = self.conf_dict['method']
        if self.method == 'score':
            self.df_src['score'] = self.df_src['score'].map(to_score)
        
        data = self.df_src.copy()
        data['type'] = 'all'
        data = pd.concat([self.df_src,data])
        self.df_src_with_all = data
        
        l = []
        range_min = self.df_src['score'].min()
        range_max = self.df_src['score'].max()
        bin_cnt = 10
        if self.method == 'score':
            bins = list(range(range_min, range_max + 1, int(round((range_max - range_min) / bin_cnt))))
            bins = bins[:bin_cnt] + [range_max + 1]
        if self.method == 'p':
           _, bins = pd.cut(self.df_src['score'], bins=10, duplicates='drop', retbins=True)
        self.bins_ = bins
        
        all_qcut = pd.qcut(self.df_src['score'], q=10, duplicates='drop', precision=0).astype(str)
        all_qcut_20 = pd.qcut(self.df_src['score'], q=self.conf_dict['cum_bins'], duplicates='drop', precision=0).astype(str)
        all_qcut_100 = pd.qcut(self.df_src['score'], q=100, duplicates='drop', precision=0).astype(str)

        ll,ll_2,ll_100 = [],[],[]
        for i in list(set(all_qcut)):
            ll.append(i.split(', ')[0][1:])
        if self.method == 'p':
            ll = [(float(x)) for x in ll]
        if self.method == 'score':
            ll = [round(float(x)) for x in ll]
        ll.sort()
        ll.extend([range_max+1])
        for i in list(set(all_qcut_20)):
            ll_2.append(i.split(', ')[0][1:])
        if self.method == 'p':
            ll_2 = [(float(x)) for x in ll_2]
        if self.method == 'score':
            ll_2 = [round(float(x)) for x in ll_2]
        ll_2.sort()
        ll_2.extend([range_max+1])
        for i in list(set(all_qcut_100)):
            ll_100.append(i.split(', ')[0][1:])
        if self.method == 'p':
            ll_100 = [(float(x)) for x in ll_100]
        if self.method == 'score':
            ll_100 = [round(float(x)) for x in ll_100]
        ll_100.sort()
        ll_100.extend([range_max+1])

        self.ll,self.ll_2,self.ll_100 = ll,ll_2,ll_100
        for splitted_type, df in self.df_src.groupby('type'):
            df['qbin'] = pd.cut(df['score'], bins=ll, precision=0, right=False).astype(str)
            df['bin'] = pd.cut(df['score'], bins=bins, precision=0, right=False).astype(str)
            l.append(df)
        self.df_src[['qbin', 'bin']] = pd.concat(l)[['qbin', 'bin']]
        self.df_grp = self.df_src.groupby('type')
  
        tr_te_begin_date = self.df_src_month[self.df_src_month['type']!='oot']['月份'].astype(int).min()
        tr_te_end_date = self.df_src_month[self.df_src_month['type']!='oot']['月份'].astype(int).max()
        oot_begin_date_,oot_end_date_ = '--','--'
        if 'oot' in set(self.df_src_month['type']):
            oot_begin_date_ = self.df_src_month[self.df_src_month['type']=='oot']['月份'].astype(int).min()
            oot_end_date_ = self.df_src_month[self.df_src_month['type']=='oot']['月份'].astype(int).max()
        self.df_detail = pd.DataFrame([[self.conf_dict['client_name'],self.df_src.shape[0],round((self.df_src.y_label==1).sum()/self.df_src.shape[0],4),
                                        tr_te_begin_date,tr_te_end_date,oot_begin_date_,oot_end_date_]],
           columns = ['合作单位','样本总量','总体坏客户比例', '训练集测试集起始日期','训练集测试集结束日期','时间外样本起始日期','时间外样本结束日期'])
        self.df_detail = self.df_detail.astype(str)

    def overview(self):
        df_overview = self.df_grp.apply(lambda df: pd.Series({
            '样本量': str(int(df.shape[0])),
            '坏客户数量': str(int(df['y_label'].sum())),
            '坏客户比例': '{:.1%}'.format(df['y_label'].mean()),
            'KS': '{:.3}'.format(get_ks(df['y_label'], df['score'])),
            'AUC': '{:.3}'.format(get_roc_auc_score(1 - df['y_label'], df['score']))
        })).T.reindex(['样本量', '坏客户数量', '坏客户比例', 'KS', 'AUC'])
        df_overview = df_overview[self.splitted_type_order]
        df_overview.index.name = ' '
        #df_overview = df_overview.reset_index()

        df= self.df_src.copy()
        df['type2'] = 'ALL'
        dic = {'样本量': str(int(df.shape[0])), '坏客户数量': str(int(df['y_label'].sum())),
               '坏客户比例': '{:.1%}'.format(int(df['y_label'].sum()) / int(df.shape[0])),
               'KS': '{:.3}'.format(get_ks(df['y_label'], df['score'])),
               'AUC': '{:.3}'.format(get_roc_auc_score(1 - df['y_label'], df['score']))}
        df_overview_all = pd.DataFrame(dic, index=['ALL']).T.reindex(['样本量', '坏客户数量', '坏客户比例', 'KS', 'AUC'])
        df_overview = pd.concat([df_overview, df_overview_all], axis=1).reset_index()
        return df_overview

    def distruibution(self, bin_col, asc):
        def single_type_distruibution(df_splitted_type):
            df_distribution = df_splitted_type.groupby(bin_col).apply(lambda df: pd.Series({
                'total_cnt': df.shape[0],
                'bad_cnt': df[df['y_label'] == 1].shape[0],
                'good_cnt': df[df['y_label'] == 0].shape[0]
            }))[['total_cnt', 'bad_cnt', 'good_cnt']]
            df_distribution = df_distribution.sort_index(ascending=asc)

            # 统计坏客户占比 = 当前箱中坏样本个数/总体坏样本个数
            df_distribution['bad_prob'] = df_distribution['bad_cnt'] / df_distribution['bad_cnt'].sum()
            # 统计好客户占比 = 当前箱中好样本个数/总体好样本个数
            df_distribution['good_prob'] = df_distribution['good_cnt'] / df_distribution['good_cnt'].sum()
            # 累计坏客户占比
            df_distribution['cum_bad_prob'] = df_distribution['bad_prob'].cumsum()
            # 累计好客户
            df_distribution['cum_good_prob'] = df_distribution['good_prob'].cumsum()
            # 区间违约率 = 当前箱中坏样本占该箱比值
            df_distribution['bad_rate'] = df_distribution['bad_cnt'] / df_distribution['total_cnt']
            # 累计通过人数
            df_distribution['cum_total_cnt'] = df_distribution['total_cnt'].cumsum()
            # 累计坏客户数
            df_distribution['cum_bad_cnt'] = df_distribution['bad_cnt'].cumsum()
            # 累计通过率 = 累计通过人数 / 总人数
            df_distribution['cum_pass'] = df_distribution['cum_total_cnt'] / df_distribution['total_cnt'].sum()
            # 累计坏人比例 = 累计通过坏人数 / 累计通过人数
            df_distribution['cum_pass_bad_rate'] = df_distribution['cum_bad_cnt'] / df_distribution['cum_total_cnt']
            df_distribution.index.name = 'score_bin'
            df_distribution = df_distribution.reset_index()

            # -------------------to str format------------------
            df_str_distruibution = df_distribution.copy()
            int_cols = ['total_cnt', 'bad_cnt', 'good_cnt', 'cum_total_cnt', 'cum_bad_cnt']
            df_str_distruibution[int_cols] = df_str_distruibution[int_cols].applymap(
                lambda v: str(int(v)) if pd.notnull(v) else '')
            rate_cols = ['bad_prob', 'good_prob', 'cum_bad_prob', 'cum_good_prob', 'bad_rate', 'cum_pass',
                         'cum_pass_bad_rate']
            df_str_distruibution[rate_cols] = df_str_distruibution[rate_cols].applymap(
                lambda v: '{:.1%}'.format(round(v, 3)) if pd.notnull(v) else '')
            df_str_distruibution = df_str_distruibution[
                ['score_bin', 'total_cnt', 'bad_cnt', 'good_cnt', 'bad_prob', 'good_prob', 'cum_bad_prob',
                 'cum_good_prob', 'bad_rate', 'cum_total_cnt', 'cum_bad_cnt', 'cum_pass', 'cum_pass_bad_rate']]
            df_str_distruibution.columns = ['评分', '组内总人数', '组内坏客户数', '组内好客户数', '坏客户占比', '好客户占比', '累计坏客户占比', '累计好客户占比',
                                            '区间违约率', '累计拒绝人数', '累计拒绝坏人数', '累计拒绝率', '累计拒绝坏人占比']
            return df_distribution, df_str_distruibution

        dict_df_distribution = {}
        dict_df_str_distribution = {}
        for splitted_type, df_splitted_type in self.df_grp:
            dict_df_distribution[splitted_type], dict_df_str_distribution[splitted_type] = single_type_distruibution(
                df_splitted_type)

        return dict_df_distribution, dict_df_str_distribution
    
    def psi(self):
        def bin_psi(x, y):

            if pd.isnull(y) or y == 0 or pd.isnull(x) or x == 0:
                return None
            else:
                return (x - y) * math.log(x / y)

        if 'train' not in self.df_grp.groups.keys():
            print('Error: failt to get psi, for train is not in splitted_types')
            return
        range_start = math.floor(df_src['score'].min() / 50) * 50
        range_stop = math.ceil(df_src['score'].max() / 50) * 50 + 1
        _, bins = pd.cut(self.df_src['score'], bins=10, duplicates='drop', retbins=True)
        l = []
        for splitted_type, df_score in self.df_grp:
            df = pd.cut(df_score['score'], bins=bins, right=False).value_counts().map(
                lambda v: v / df_score.shape[0]).to_frame('pct')
            df.index.name = 'bin'
            df.index = df.index.astype(str)
            df = df.reset_index()
            df['type'] = splitted_type
            l.append(df)

        df_psi_detail = pd.concat(l, ignore_index=True).pivot_table(index='bin', columns='type', values='pct')
        df_psi_detail.columns = [s + '_pct' for s in df_psi_detail.columns.format()]
        df_psi_detail = df_psi_detail.reset_index()
        for splitted_type in filter(lambda s: s != 'train', self.df_grp.groups.keys()):
            df_psi_detail['train_{}_psi'.format(splitted_type)] = df_psi_detail.apply(
                lambda r: bin_psi(r['train_pct'], r[splitted_type + '_pct']), axis=1)

        df_psi_detail_sum = df_psi_detail.drop(labels='bin', axis=1).sum()
        df_psi_detail_sum['bin'] = 'sum'
        df_psi_detail_sum = df_psi_detail_sum.to_frame().T
        df_psi_detail = pd.concat([df_psi_detail, df_psi_detail_sum], ignore_index=True)

        # format
        pct_cols = [x + '_pct' for x in self.splitted_type_order]
        psi_cols = ['train_{}_psi'.format(x) for x in filter(lambda s: s != 'train', self.splitted_type_order)]
        df_psi_detail[pct_cols] = df_psi_detail[pct_cols].applymap(
            lambda v: '{:.1%}'.format(v) if pd.notnull(v) else '')
        df_psi_detail[psi_cols] = df_psi_detail[psi_cols].applymap(
            lambda v: '{}'.format(round(v, 3)) if pd.notnull(v) else '')
        df_psi_detail = df_psi_detail[['bin'] + pct_cols + psi_cols]
        df_psi_detail.rename(columns={'bin': '评分','train_pct':'train样本量占比','test_pct':'test样本量占比'}, inplace=True)
        return df_psi_detail

    def statis(self,asc):
        self.df_overview = self.overview()
        self.df_psi = self.psi()
        self.dict_df_distribution_equidistance, self.dict_df_str_distribution_equidistance = self.distruibution('bin', asc)
        self.dict_df_distribution_equifrequency, self.dict_df_str_distribution_equifrequency = self.distruibution('qbin', asc)

    def plot_tot_describe(self, output_dir):
        font = FontProperties(fname=r"simsun.ttc", size=14)
        for splitted_type, df_splitted_type in self.df_src_with_all.groupby('type'):
            plt.figure(figsize=(6,4))
            sns.distplot( df_splitted_type[df_splitted_type['type']==splitted_type]["score"] , color='#ff8080',bins=20,kde_kws={ "lw": 2.5, 'linestyle':'--'})
            tem = df_splitted_type[df_splitted_type['type']==splitted_type]["score"].describe().reset_index()
            table_ = plt.table(cellText = [[round(x,4)] for x in tem['score'].tolist()],
                      colWidths=[0.1]*1,rowLabels = tem['index'].tolist(),loc='right')
            table_.set_fontsize(15)
            table_.scale(1.9,2.265)
            plt.title('{}评分分布'.format(splitted_type), fontproperties=font)
            plt.savefig(output_dir + '{}评分分布'.format(splitted_type), bbox_inches='tight')
        color_list = ['#d65a31','#40bfc1','#b7e778','#5edfff']
        plt.figure(figsize=(6,4))
        a = -1
        for splitted_type, df_splitted_type in self.df_src_with_all.groupby('type'):
            a = a+1
            sns.kdeplot(df_splitted_type[df_splitted_type['type']==splitted_type]["score"] ,shade_lowest=False,color=color_list[a],
                        label = splitted_type)
        plt.title('各样本的评分分布对比', fontproperties=font)
        plt.savefig(output_dir + '各样本的评分分布对比', bbox_inches='tight')
        return

    def cumsum_(self,split_type,bins):
        if 'oot' in self.df_src.type:
            self.df_src_cum = self.df_src[self.df_src.type=='oot'].reset_index(drop = True)
        else:
            self.df_src_cum = self.df_src[self.df_src.type==split_type].reset_index(drop = True)
        #确定十分位点
#         bins = pd.cut(self.df_src_cum['score'],bins =20,retbins =True)[1]# 等宽分箱
        #bins = pd.qcut(self.df_src_cum['score'], q=10, duplicates='drop', precision=0).astype(str)# 等频
#         bins = [int(i) for i in bins]
#         self.ll_2[0] = 299
#         self.ll_2[-1] =900
        self.cg_dict ={'row_0':'评分区间','index':'评分区间',1:'坏客户',0:'好客户','group%':'区间占比','good_goodall%':'好客户占比','bad_badall%':'坏客户占比','bad%':'区间违约率',
                  'cum_bad%':'累积坏客户占比','cum_good%':'累积好客户占比','cum_reject':'累积拒绝人数','cum_bad':'累积拒绝坏人数',
                  'cum_reject%':'拒绝率','cum_bad%_reject':'拒绝坏人占比','lift':'Lift'}
        # 计算十分箱
        
        bins_10 = crosstab_bins(X = self.df_src_cum['score'],y = self.df_src_cum['y_label'],bins =bins,reverse_ = 0).reset_index()
        bins_10.rename(columns =self.cg_dict,inplace =True)
        bins_10['评分区间'] = bins_10['评分区间'].astype(str)
        bins_10_r = crosstab_bins(X = self.df_src_cum['score'],y = self.df_src_cum['y_label'],bins =bins,reverse_ = 1).reset_index()
        bins_10_r.rename(columns =self.cg_dict,inplace =True)
        bins_10_r['评分区间'] = bins_10_r['评分区间'].astype(str)
        tot_bad_rate = (self.df_src_cum.y_label==1).sum()/self.df_src_cum.shape[0]
        return bins_10,bins_10_r,tot_bad_rate
    
    def plot_cumsum(self, output_dir,data,title,splite_type):
        data['拒绝坏人占比'] = data['拒绝坏人占比'].astype(float)
        data['累积坏客户占比'] = data['累积坏客户占比'].astype(float)
        width = 0.4
        xticks = np.arange(data.shape[0])
        fig = plt.figure(figsize=(8, 3))
        ax1 = fig.add_subplot(111)
        p1 = ax1.bar(xticks, data['拒绝率'].astype(float), width, color='#9ddcdc')
        p4 = ax1.plot(xticks, data['累积坏客户占比'],'*-', color='#43ab92')
        t_c = ['{:.1%}'.format(round(x, 3)) if pd.notnull(x) else '' for x in data['累积坏客户占比']] 
        for a, b in zip(xticks, [round(float(x),2) for x in data['拒绝率']]):
            plt.text(a, b, b, ha='center', va='bottom', fontsize=9,color ='#160f30')
        for a, b ,c in zip(xticks, data['累积坏客户占比'],t_c):
            plt.text(a, b, c, ha='center', va='bottom', fontsize=9.5,color ='#160f30')
            
        ax1.set_ylabel(u'累积比例')
        ax1.set_xlabel(u'区间')
        ax1.set_title(splite_type+title)
        ax1.set_xticks(xticks)
        ax1.set_xticklabels(data['评分区间'].astype(str), rotation=80)
        ax2 = ax1.twinx()
#         p2 = ax2.plot(xticks, data['拒绝率'].astype(float),'o-', color='#207561')
        p3 = ax2.plot(xticks, data['拒绝坏人占比'],'o-', color='#d65a31')
        t_b = ['{:.1%}'.format(round(x, 3)) if pd.notnull(x) else '' for x in data['拒绝坏人占比']] 
        for a, b ,c in zip(xticks,data['拒绝坏人占比'], t_b):
            plt.text(a, b, c, ha='center', va='bottom', fontsize=9.5,color ='#160f30')
#         ax2.plot([self.tot_bad_rate]*20,"r--")
        ax2.set_ylabel(u'违约率')
        ax2.set_ylim(0, (data['拒绝坏人占比']).max()*1.2)
        plt.legend((p1[0],p3[0],p4[0]), ('累积比例','累积的违约率','累积坏客户占比'))
        plt.savefig(output_dir + splite_type + title, bbox_inches='tight')
        return
    
    def advice_(self,data1,data2):
        tem = data1[data1['拒绝坏人占比']>=self.conf_dict['advice_list']['reject_bad_rate']].iloc[-1:,:]
        tem_b1 = float(tem['评分区间'].values[0].split(', ')[1][:-1])
        tem = data2[data2['拒绝坏人占比']<=self.conf_dict['advice_list']['accept_bad_rate']].iloc[-1:,:]
        tem_b2 = float(tem['评分区间'].values[0].split(', ')[0][1:])
        print(tem_b1,tem_b2)
        self.bins_3 = crosstab_bins(X = self.df_src_cum['score'].astype(float),y = self.df_src_cum['y_label'],bins =[299,tem_b1,tem_b2,900],reverse_ = 0).reset_index()
        self.bins_3.rename(columns =self.cg_dict,inplace =True)
        self.bins_3['建议']=['拒绝','其他','通过']
        self.bins_3 = self.bins_3[['建议']+list(self.cg_dict.values())[1:]+['评分区间']]
        self.df_src_cum['suggestion_1'] = ['通过' if x>tem_b2 else '拒绝' if x<=tem_b1 else '复审'  for x in  self.df_src_cum['score']]
        del self.bins_3['拒绝率']

        self.cross_1 = pd.crosstab(self.df_src_cum['suggestion_1'],self.df_src_cum['y_label'],margins=True).reset_index()
        self.cross_1.columns = [' ','好','坏','ALL']
        self.cross_1['违约率'] = list(map(lambda x,y:x/y ,self.cross_1['坏'],self.cross_1['ALL']))
        self.cross_1['违约率'] = self.cross_1['违约率'].map(lambda v: '{:.1%}'.format(round(v, 3)) if pd.notnull(v) else '' )
        self.cross_1 = self.cross_1.astype(str)
        
        self.tem_b1,self.tem_b2 = tem_b1,tem_b2
        return
    
    def plot_advice(self, output_dir,data,title):
        width = 0.35
        xticks = np.arange(data.shape[0])
        data = data.sort_values(by = '建议').fillna(0)
        fig = plt.figure(figsize=(8, 4))
        ax1 = fig.add_subplot(111)
        p1 = ax1.bar(xticks, data['好客户占比'], width, color='#9ddcdc')
        p2 = ax1.bar(xticks+0.35, data['坏客户占比'], width, color='#ffb6b9')
        ax1.set_ylabel(u'用户占比')
        ax1.set_title('策略建议')
        ax1.set_xticks(xticks)
        ax1.set_xticklabels(data['建议'].astype(str), rotation=0)
        ax2 = ax1.twinx()
        p3 = ax2.plot(xticks, data['区间违约率'],'o-', color='#d65a31')
        p4 = ax2.plot(xticks, data['区间占比'],'*-', linestyle='--',color='#00b8a9')
        for a, b in zip(xticks, [round(x,4) for x in data['区间占比']]):
            plt.text(a, b, b, ha='center', va='bottom', fontsize=12,color ='#3a4750')
        ax2.set_ylabel(u'区间违约率')
        ax2.set_ylim(0, 0.8)
        plt.legend((p1[0], p2[0],p3[0],p4[0]), ('好客户占比', '坏客户占比','区间违约率','区间占比'))
        plt.savefig(output_dir + title, bbox_inches='tight')
        return
   
    @staticmethod 
    def g_num(group): return (group==0).sum()
    @staticmethod
    def b_num(group): return (group==1).sum()
    @staticmethod 
    def g_r(group): return (group==1).sum()/group.shape[0]
    
    def df_monthly(self, output_dir,ix = '月份'):
        self.ix = ix
        self.monthly = self.df_src_month.groupby(self.df_src_month[ix],as_index=True).agg({'y_label':[self.g_num,self.b_num,'count',self.g_r]}).reset_index()
        self.monthly.columns = [ix,'good人数','bad人数','总人数','bad_rate']
        self.monthly_tr = self.df_src_month[self.df_src_month['type']=='train'].groupby(self.df_src_month[ix],as_index=True).agg({'y_label':[self.g_num,self.b_num,'count',self.g_r]}).reset_index()
        self.monthly_tr.columns = [ix,'tr_good人数','tr_bad人数','tr_总人数','tr_bad_rate']
        self.monthly_te = self.df_src_month[self.df_src_month['type']=='test'].groupby(self.df_src_month[ix],as_index=True).agg({'y_label':[self.g_num,self.b_num,'count',self.g_r]}).reset_index()
        self.monthly_te.columns = [ix,'te_good人数','te_bad人数','te_总人数','te_bad_rate']
        self.monthly = self.monthly.merge(self.monthly_tr,on = ix,how = 'left')
        self.monthly = self.monthly.merge(self.monthly_te,on = ix,how = 'left')
        self.monthly = self.monthly.sort_values(by =ix).fillna(0)
        width = 0.4
        xticks = np.arange(self.monthly.shape[0])
        fig = plt.figure(figsize=(8, 3))
        ax1 = fig.add_subplot(111)
        p1 = ax1.bar(xticks, self.monthly['good人数'].astype(int), width, color='#9ddcdc')
        p2 = ax1.bar(xticks, self.monthly['bad人数'].astype(int), width,bottom=self.monthly['good人数'], color='#ffb6b9')
        ax1.set_ylabel(u'用户数')
        ax1.set_xlabel(u'时间段')
        ax1.set_title('按{}统计的客户分布'.format(self.ix))
        ax1.set_xticks(xticks)
        ax1.set_xticklabels(self.monthly[ix].astype(str), rotation=80)
        ax2 = ax1.twinx()
        p3 = ax2.plot(xticks, self.monthly['bad_rate'],'o-', color='#d65a31')
        for a, b in zip(xticks, [round(x,2) for x in self.monthly['bad_rate']]):
            plt.text(a, b, b, ha='center', va='bottom', fontsize=12,color ='#d65a31')
        ax2.set_ylabel(u'bad_rate')
        ax2.set_ylim(0, self.monthly['bad_rate'].max()*1.2)
        plt.legend((p1[0], p2[0],p3[0]), ('好客户数量', '坏客户数量','坏客户比例'))
        plt.savefig(output_dir +'按{}统计的客户分布'.format(self.ix), bbox_inches='tight')
        return
        
    def to_excel(self, output_path):
        writer = pd.ExcelWriter(output_path)
        self.df_overview.to_excel(writer, sheet_name='overview', index=False)
        self.df_psi.to_excel(writer, sheet_name='psi', index=False)
        for k, df in self.dict_df_str_distribution_equidistance.items():
            df.to_excel(writer, sheet_name='{}_distrubution_eq-dst'.format(k), index=False)
        for k, df in self.dict_df_str_distribution_equifrequency.items():
            df.to_excel(writer, sheet_name='{}_distrubution_eq-frq'.format(k), index=False)
        writer.save()

    def plot_roc(self, output_dir):
        for splitted_type, df_splitted_type in self.df_src_with_all.groupby('type'):
            title = splitted_type+' ROC'
            if self.method == 'p':
                plot_roc(df_splitted_type['y_label'],df_splitted_type['score'], output_path=output_dir+title,title=title)
            if self.method == 'score':
                plot_roc(1-df_splitted_type['y_label'],df_splitted_type['score'], output_path=output_dir+title,title=title)           
        return
    
    def plot_roc2(self, output_dir):
        plt.figure(figsize=(4,4))
        lw = 1
        for splitted_type, df_splitted_type in self.df_src_with_all.groupby('type'):
            if self.method == 'p':
                fpr, tpr, thresholds = metrics.roc_curve(df_splitted_type['y_label'], df_splitted_type['score'])
            if self.method == 'score':
                fpr, tpr, thresholds = metrics.roc_curve(1-df_splitted_type['y_label'], df_splitted_type['score'])
            plt.plot(fpr, tpr,lw=lw, label=splitted_type+'_curve (area = %0.4f)' % auc(fpr,tpr))

        plt.plot([0, 1], [0, 1], color='navy', lw=lw, linestyle='--')
        plt.xlim([0.0, 1.0])
        plt.ylim([0.0, 1.05])
        plt.xlabel('False Positive Rate')
        plt.ylabel('True Positive Rate')
        plt.title('ALL_TR_TE_OOT ROC')
        plt.legend(loc="lower right")
        plt.savefig(output_dir + 'ALL_TR_TE_OOT ROC', bbox_inches='tight')
    
    def plot_ks_(self, output_dir):
        for splitted_type, df_splitted_type in self.df_src_with_all.groupby('type'):
            title = splitted_type+' KS'
            if self.method == 'p':
                plot_ks(df_splitted_type['y_label'],df_splitted_type['score'], output_path=output_dir+title,title=title)
            if self.method == 'score':
                plot_ks(df_splitted_type['y_label'],-df_splitted_type['score'], output_path=output_dir+title,title=title)                
        return
    
    def plot_lift_(self,output_dir):
        for splitted_type, df_splitted_type in self.df_src_with_all.groupby('type'):
            title = splitted_type+' Lift'
            tem_qcut = pd.cut(df_splitted_type['score'], bins=self.ll_2, precision=0, right=False).astype(str)
            asc = False if self.method == 'p' else True
            plot_lift(df_splitted_type['y_label'],
#                       pd.qcut(df_splitted_type['score'], q=10, duplicates='drop', precision=0),
                        tem_qcut,output_path=output_dir+title,title=title, asc=asc)
        
        fig = plt.figure(figsize=(8, 3))
        ax1 = fig.add_subplot(111)
        color_list = ['#d65a31','#40bfc1','#589167','#5edfff']
        markers_list = ['o-','x-','^-','+-']
        a = 0
        for splitted_type, df_splitted_type in self.df_src_with_all.groupby('type'):
            tem_qcut = pd.cut(df_splitted_type['score'], bins=self.ll_2, precision=0, right=False).astype(str)
            data = pd.concat([df_splitted_type['y_label'],tem_qcut], axis=1)
            data.columns = ['y_true', 'bin']
            b_t = (data['y_true']==1).sum()
            data = data.groupby('bin').apply(lambda df: pd.Series({'cnt': df.shape[0], 'bad_cnt': df['y_true'].sum()}))
            if self.method == 'p':
                data = data.sort_index(ascending=False)
            if self.method == 'score':
                data = data.sort_index(ascending=True)
            data = data.cumsum()
            data['bad_cnt_rate']= data['bad_cnt']/b_t
            data['cnt_rate']= data['cnt']/df_splitted_type.shape[0]
            data['lift'] = data['bad_cnt_rate']/data['cnt_rate']
            data = data.reset_index()
            x = np.arange(data.shape[0])
            p1 = ax1.plot(x, [round(x,2) for x in data['lift'].astype(float)],markers_list[a],color =color_list[a], label = splitted_type)
            a = a+1
            ax1.set_ylabel(u'提升度')
            ax1.set_xlabel(u'等频分箱')
            ax1.set_title('各样本的提升度对比')
            ax1.set_xticks(x)
            ax1.set_xticklabels(list(data['bin']), rotation=80)
            plt.xticks(rotation=80)
            plt.legend() # 显示图例
        plt.savefig(output_dir+'各样本的提升度对比', bbox_inches='tight')
        return

    def plot_bins_(self, output_dir):
        for splitted_type, df_splitted_type in self.df_src_with_all.groupby('type'):
            title = u'{}分箱统计'.format(splitted_type)
            plot_bins(pd.qcut(df_splitted_type['score'], q=10, duplicates='drop', precision=5),
                      df_splitted_type['y_label'],
                      output_path=output_dir+title,title=title)
        return

    def plot_distribution2(self, dict_df_distribution, distribution_type, output_dir):
        for k, df_distribution in dict_df_distribution.items():
            title = '{}{}分布'.format(k, distribution_type)
            x = np.arange(df_distribution.shape[0])
            y1 = list(df_distribution['bad_prob'])
            y2 = list(df_distribution['good_prob'])
            y3 = list(df_distribution['bad_rate'])
            font = FontProperties(fname=r"simsun.ttc", size=14)
            width = 0.4
            fig = plt.figure(figsize=(8, 3))
            ax1 = fig.add_subplot(111)
            p1 = ax1.bar(x, y1, width, color='#ffb6b9')
            p2 = ax1.bar(x+width, y2, width, color='#9ddcdc')
            ax1.set_ylabel(u'用户数', fontproperties=font)
            ax1.set_xlabel(u'分箱编号', fontproperties=font)
            ax1.set_title(title, fontproperties=font)
            ax1.set_xticks(x)
            ax1.set_xticklabels(list(df_distribution['score_bin']), rotation=90)
            
#             ax2 = ax1.twinx()
#             p3 = ax2.plot(x, y1,'o-', color='#d65a31')
#             p4 = ax2.plot(x+width, y2,'*-', color='#207561')
#             ax2.set_yy_label(u'区间违约率', fontproperties=font)
            for a, b in zip(x, [round(x,2) for x in y1]):
                plt.text(a, b, b, ha='center', va='bottom', fontsize=10,color ='#3c4245')
            for a, b in zip(x, [round(x,2) for x in y2]):
                plt.text(a+width, b, b, ha='center', va='bottom', fontsize=10,color ='#3c4245')
            
            plt.legend((p1[0], p2[0]), ('坏客户占总坏客户比例', '好客户占总好客户比例'),loc = 'upper right')
            plt.savefig(output_dir + title, bbox_inches='tight')
        return
    
    def plot(self, output_dir):
        self.df_monthly(output_dir,ix = self.conf_dict['statistics_by_month'])
        self.plot_roc(output_dir)
        self.plot_roc2(output_dir)
        self.plot_tot_describe(output_dir)
        self.plot_ks_(output_dir)
        self.plot_lift_(output_dir)
        self.plot_bins_(output_dir)
        self.plot_distribution2(self.dict_df_distribution_equidistance, '等宽', output_dir)
        self.bins_10,self.bins_10_r,self.tot_bad_rate_10 = self.cumsum_(split_type = 'test',bins = self.ll_2)
        self.bins_100,self.bins_100_r,self.tot_bad_rate_100 = self.cumsum_(split_type = 'test',bins = self.ll_100)
#         print(self.bins_100.iloc[:10,:5])
        self.plot_cumsum(output_dir,self.bins_10,'向上累积统计','test')
        self.plot_cumsum(output_dir,self.bins_100.iloc[:20,:],'尾部向上累积统计','test')
        self.plot_cumsum(output_dir,self.bins_10_r,'向下累积统计','test')
        self.plot_cumsum(output_dir,self.bins_100_r.iloc[:20,:],'尾部向下累积统计','test')
        if self.conf_dict['is_output_advice']=='y':
            try:
                self.advice_(self.bins_10,self.bins_10_r)
                self.plot_advice(output_dir,self.bins_3,'advice_1')
            except:
                print('策略设置不成功：拒绝坏账过高或者通过坏账过低，或拒绝加通过总比例大于1，请查看累积分布重新选择')
        return

    @staticmethod
    def document_add_df(document, df, header=True):
        header_row_cnt = 1 if header else 0
        table = document.add_table(rows=df.shape[0] + header_row_cnt, cols=df.shape[1], style='Light List Accent 5')
        table.autofit = True
        if header:
            header_cells = table.rows[0].cells
            for i, c in enumerate(df.columns.format()):
                run_ = header_cells[i].paragraphs[0].add_run(c)              
                run_.font.name = '宋体'
                run_.font.size = Pt(10)
                header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for row_idx in range(0, df.shape[0]):
            for col_idx in range(0, df.shape[1]):
#                 cell = table.cell(row_idx + header_row_cnt, col_idx)
#                 cell.text = df.iloc[row_idx, col_idx]
                th(table,row_idx+header_row_cnt,col_idx,df.iloc[row_idx,col_idx])
        document.add_paragraph()
        return

    def to_report(self, graph_dir, output_path, title,cbid):
        document = Document('../data/report/标准模板.docx')
        section = document.sections[-1]
        (section.top_margin,
         section.bottom_margin,
         section.left_margin,
         section.right_margin) = tuple([Cm(1)] * 4)

        font = document.styles['Normal'].font
        font = '宋体'
        picture_size = Cm(9)
     
        run = document.add_heading('',level=0).add_run('AKULAKU：{}建模报告'.format(self.conf_dict['client_id_name']))
        run.font.name=u'华文中宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')
        for blank in range(5):
            document.add_paragraph('')
        page1 =  [['制作部门', 'AKULAKU风控建模'],
                 ['使用模型', self.conf_dict['algorithm']],
                 ['模型版本', self.conf_dict['model_version']],
                 ['制作单位', 'AKULAKU'],
                 ['制作时间', datetime.now().strftime( '%Y-%m-%d %H:%M:%S %f' )[:10]]]
        for blank in range(1):
            document.add_paragraph('')
            
        tb=document.add_table(rows=4,cols=2)
        tb.add_row()
        for row in range(len(page1)):
            for col in range(2):
                tb.cell(row,col).width=1
                tb.cell(row,col).text=page1[row][col]
                tb.cell(row,col).width=Cm(6)

        tb.style='Table Grid'
        tb.autofit=True
        tb.alignment=WD_TABLE_ALIGNMENT.CENTER
        document.add_page_break()
        
        UserStyle1 = document.styles.add_style('UserStyle1', 1)
        UserStyle1.font.size = Pt(10)
        UserStyle1.font.color.rgb = RGBColor(33, 23, 23)
        UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        UserStyle1.font.name = '微软雅黑'
        UserStyle1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        UserStyle2 = document.styles.add_style('UserStyle2', 1)
        UserStyle2.font.size = Pt(12)
        UserStyle2.font.color.rgb = RGBColor(33, 23, 23)
        UserStyle2.font.name = '宋体'
        UserStyle2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
              
        UserStyle3 = document.styles.add_style('UserStyle3', 1)
        UserStyle3.font.size = Pt(8)
        UserStyle3.font.color.rgb = RGBColor(192,192,192)
        UserStyle3.font.name = '微软雅黑'
        UserStyle3._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        document.add_heading('一、建模样本', 1)
        document.add_paragraph('表1 样本概况表',style = UserStyle1)
        ML_External_Report.document_add_df(document, self.df_detail.astype(str))
        rate_cols = ['bad_rate','tr_bad_rate','te_bad_rate']
        self.monthly[rate_cols] = self.monthly[rate_cols].applymap(lambda v: '{:.1%}'.format(round(v, 3)) if pd.notnull(v) else '' )
        document.add_paragraph('表2 按{}统计分布表'.format(self.conf_dict['statistics_by_month']),style = UserStyle1)
        self.monthly.columns = [self.conf_dict['statistics_by_month'],'好客户数','坏客户数','总数','总违约率','训练集好客户数','训练集坏客户数','训练集总数','训练集违约率','测试集好客户数','测试集坏客户数','测试集总数','测试集违约率']
        ML_External_Report.document_add_df(document, self.monthly.astype(str))
        document.add_picture(graph_dir + '按{}统计的客户分布.png'.format(self.conf_dict['statistics_by_month']))
        document.add_page_break()
        
        document.add_heading('二、效果概况', 1)
        document.add_paragraph('表3 模型效果概况表',style = UserStyle1)
        ML_External_Report.document_add_df(document, self.df_overview.astype(str))
        for splitted_type in self.splitted_type_order+['all']:
            document.add_picture(graph_dir + '{}评分分布.png'.format(splitted_type))
        document.add_picture(graph_dir + '各样本的评分分布对比.png')
        document.add_page_break()
        
        document.add_heading('三、ROC曲线', 1)
        document.add_paragraph("ROC（Eceiver operating characteristic curve）：以FPR为横轴，TPR为纵轴，在不同阈值下计算FPR和TPR的值画出的图形。ROC曲线同对角线形成的面积越大说明模型排序能力越好。",style = UserStyle2)
        p = document.add_paragraph()
        for splitted_type in self.splitted_type_order+['all']:
            p.add_run().add_picture(graph_dir + '{} ROC.png'.format(splitted_type), width=picture_size)
        p.add_run().add_picture(graph_dir + 'all_tr_te_oot ROC.png'.format(splitted_type), width=picture_size)
        document.add_page_break()
        
        document.add_heading('四、KS曲线', 1)
        document.add_paragraph("KS（Kolmogorov-Smirnov）：通过衡量好坏样本累计分布之间的差值，来评估模型的风险区分能力。KS值越大，区分度越强。",UserStyle2)
        p = document.add_paragraph()
        for splitted_type in self.splitted_type_order+['all']:
            p.add_run().add_picture(graph_dir + '{} KS.png'.format(splitted_type), width=picture_size)
        document.add_page_break()
        
        document.add_heading('五、Lift曲线', 1)
        document.add_paragraph("Lift：提升度，表示与不利用模型相比，模型的预测能力提升多少，lift(提升指数)越大，模型效果越好。",UserStyle2)
        document.add_paragraph("将评分等频分为20份，坏客户占总坏客户的累积比例，比样本数占总体样本数的累积比例，得提升度Lift，即该模型抓取坏客户的能力是随机选择的多少倍。",UserStyle2)
        p = document.add_paragraph()
        for splitted_type in self.splitted_type_order+['all']:
            p.add_run().add_picture(graph_dir + '{} Lift.png'.format(splitted_type))
        p.add_run().add_picture(graph_dir + '各样本的提升度对比.png')
        document.add_page_break()
        
        document.add_heading('六、等频分箱', 1)
        document.add_paragraph("将训练集、测试集、总体分数均分为数量相同的10个区间，每个区间对应的好客户、坏客户数量及区间违约率。",UserStyle2)
        p = document.add_paragraph()
        for splitted_type in self.splitted_type_order+['all']:
            p.add_run().add_picture(graph_dir + '{}分箱统计.png'.format(splitted_type), width=picture_size)
        document.add_page_break()
        
        document.add_heading('七、分数分布', 1)
        document.add_paragraph("将训练集、测试集分数分为分数间距相同的10个区间，每个区间对应的好客户、坏客户占比及区间违约率。",UserStyle2)
        a = 4
        for i, splitted_type in enumerate(self.splitted_type_order):
            document.add_heading('{}.{}'.format(i + 1, splitted_type), 2)
            document.add_paragraph('表{} 等宽分箱数据分布表'.format(a),style = UserStyle1)
            ML_External_Report.document_add_df(document, self.dict_df_str_distribution_equidistance[splitted_type])
            document.add_picture(graph_dir + '{}等宽分布.png'.format(splitted_type))
            a = a+1
        document.add_page_break()
        
        document.add_heading('八、累积分数分布', 1)
        document.add_paragraph("测试集等频分为20个区间，进行向上累积、向下累积。可通过向上累积选择拒绝阈值，通过向下累积选择通过阈值。",UserStyle2)
        document.add_heading('1. 向上累积分布', 2)
        document.add_picture(graph_dir + 'test向上累积统计.png')
        document.add_picture(graph_dir + 'test尾部向上累积统计.png')
        document.add_heading('2. 向下累积分布', 2)
        document.add_picture(graph_dir + 'test向下累积统计.png')
        document.add_picture(graph_dir + 'test尾部向下累积统计.png')
        document.add_page_break()

        document.add_heading('九、PSI', 1)
        document.add_paragraph('PSI（Population stability index）：可衡量测试样本及模型开发样本评分的的分布差异，检验变量的稳定性。psi值越小，变量分布差异越小，越稳定。',UserStyle2)
        document.add_paragraph('表6 PSI计算表',style = UserStyle1)
        ML_External_Report.document_add_df(document, self.df_psi)
        document.add_page_break()
        
        document.add_heading('十、策略建议', 1)
        if self.conf_dict['is_output_advice']=='y':
            document.add_paragraph("若有OOT样本外数据，则阈值在样本外选择，若无，则阈值在TEST测试集上选择。",UserStyle2)
            document.add_heading('1. 建议', 2)
            document.add_paragraph('从坏客户角度选择：按违约率小于{}来选通过阈值，按违约率大于{}来选拒绝阈值。'.format(self.conf_dict['advice_list']['accept_bad_rate'],self.conf_dict['advice_list']['reject_bad_rate']),UserStyle2)
            document.add_paragraph('通过阈值：{}（大于等于{}建议通过）'.format(self.tem_b2,self.tem_b2),UserStyle2)
            document.add_paragraph('拒绝阈值：{}（小于{}建议拒绝）'.format(self.tem_b1,self.tem_b1),UserStyle2)
            document.add_paragraph('其他：建议加入其他策略（规则、模型或人工审核）',UserStyle2)
            rate_cols = ['区间占比','好客户占比','坏客户占比','区间违约率','累积坏客户占比','累积好客户占比','拒绝坏人占比','Lift']
            self.bins_3[rate_cols] = self.bins_3[rate_cols].applymap(lambda v: '{:.1%}'.format(round(v, 3)) if pd.notnull(v) else '' )
            self.bins_3 = self.bins_3.astype(str)
            self.bins_3 = self.bins_3.drop(['累积坏客户占比','累积好客户占比','累积拒绝人数','累积拒绝坏人数','拒绝坏人占比','Lift'],axis = 1)
            document.add_paragraph('表7 策略效果表',style = UserStyle1)
            ML_External_Report.document_add_df(document, self.bins_3.iloc[:,:-1])
            document.add_picture(graph_dir + 'advice_1.png')
            document.add_page_break()
            document.add_paragraph('表8 预测结果和实际结果交叉表',style = UserStyle1)
            ML_External_Report.document_add_df(document, self.cross_1)

            document.add_heading('2. 说明', 2)
            document.add_paragraph('以上策略建议仅供参考。请贵公司结合其他数据或自有评分，按照风险容忍度、通过率要求等自身业务需求定制阈值。',UserStyle2)

        for blank in range(4):
            document.add_paragraph('')

        output_path = output_path.split('_')[0] + '_' + self.method + '_report.docx'  
        document.save(output_path)

def read_from_file(path):
    stdin = codecs.open(path, 'rb', 'utf8')
    content = stdin.read()
    stdin.close()
    return content

def load_json(path):
    content = read_from_file(path)
    return json.loads(content)

def th(table_,x,y,content):
    run = table_.cell(x,y).paragraphs[0].add_run(content)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    table_.cell(x,y).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
conf_dict = load_json('report_conf.json')
cbid = conf_dict['client_id_name']
data_dir = conf_dict['data_dir']
score = conf_dict['score_outsideReport']
asc = bool(conf_dict['asc'])

if __name__ == '__main__':
    df_src, feature_cols, score_cols = get_df_src(cbid, path_df_id=data_dir + u'{}_split.pkl '.format(cbid),
                                                  score_dir=data_dir + 'score/', scores=[score])
    df_src = df_src[df_src['type'].notnull() & df_src[score].notnull()]
    df_src['score'] = df_src[score]

    df_src.rename(columns={'y_y_label': 'y_label'}, inplace=True)
    df_src = df_src.reset_index()
    df_src_month = df_src[['uid', 'type', 'score', 'triger_date','y_label']]
    
    df_src_month['月份'] =  [str(x).replace("/",'-') for x in df_src_month['triger_date']]
    try:
        df_src_month['月份'] = [str(datetime.strftime(datetime.strptime(x,"%Y-%m-%d"),"%Y%m%d"))[:6] for x in df_src_month['月份']]
    except:
        try:
            df_src_month['月份'] = [str(datetime.strftime(datetime.strptime(x,"%Y-%m"),"%Y%m"))[:6] for x in df_src_month['月份']]
        except:
            df_src_month['月份'] = [str(datetime.strftime(datetime.strptime(x,"%Y%m%d"),"%Y%m%d"))[:6] for x in df_src_month['月份']]
    df_src_month['季度'] =  ['{}Q{}'.format(x[:4],(int(x[-2:])-1)//3+1) for x in df_src_month['月份']]
    df_src = df_src[['uid', 'type', 'score', 'y_label']]

    excel_path = data_dir + 'statis/{}_{}_external_statis.xlsx'.format(cbid, score)
    graph_dir = data_dir + 'statis/graph/{}_{}_external/'.format(cbid, score)
    os.makedirs(graph_dir, exist_ok=True)

    title = '{}_{}_report'.format(cbid, score)
    report_path = data_dir + 'report/{}.docx'.format(title)

    ml_external_report = ML_External_Report(df_src,df_src_month,conf_dict)
    ml_external_report.statis(asc)
    ml_external_report.to_excel(excel_path)
    ml_external_report.plot(graph_dir)
    ml_external_report.to_report(graph_dir, report_path, title,cbid)
    print('finish..')